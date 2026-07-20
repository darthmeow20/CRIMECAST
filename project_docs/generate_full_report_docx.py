# -*- coding: utf-8 -*-
"""
Generate CRIMECAST full project report (.docx) for hard-copy submission.
Index structure matches project_docs screenshot (college report format).

Run from project root OR from project_docs:
  py -3 project_docs/generate_full_report_docx.py
"""
from __future__ import annotations

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Inches, Pt, Cm, RGBColor
except ImportError:
    print("Install python-docx:  py -3 -m pip install python-docx")
    sys.exit(1)

ROOT = Path(__file__).resolve().parent.parent
OUT = Path(__file__).resolve().parent / "CRIMECAST_FULL_PROJECT_REPORT.docx"
# Prefer NEW regenerated figures (project_docs/figures) over legacy model_outputs/figures
FIGS_NEW = Path(__file__).resolve().parent / "figures" / "results"
SCR_NEW = Path(__file__).resolve().parent / "figures" / "screenshots"
FIGS_LEGACY = ROOT / "model_outputs" / "figures"
DIAG = ROOT / "reports" / "diagrams"
SHOT = Path(__file__).resolve().parent
LEGACY_SHOT = ROOT / "report_materials" / "screenshots"


def pick_fig(*candidates: str) -> Path | None:
    """Prefer project_docs/figures/results, then legacy model_outputs/figures."""
    for name in candidates:
        for base in (FIGS_NEW, FIGS_LEGACY):
            p = base / name
            if p.exists():
                return p
    # any png matching prefix
    for name in candidates:
        stem = name.split(".")[0]
        for base in (FIGS_NEW, FIGS_LEGACY):
            if not base.exists():
                continue
            hits = sorted(base.glob(f"*{stem}*")) + sorted(base.glob(f"{stem}*"))
            if hits:
                return hits[0]
    return None


def pick_shot(*names: str) -> Path | None:
    for name in names:
        p = SCR_NEW / name
        if p.exists():
            return p
    for p in sorted(SCR_NEW.glob("shot_*.png")) if SCR_NEW.exists() else []:
        return p
    return None


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def p(doc, text, *, size=12, bold=False, italic=False, align="left", space_after=8, space_before=0):
    para = doc.add_paragraph()
    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == "right":
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    pf = para.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return para


def h1(doc, text):
    para = doc.add_heading(text, level=1)
    for run in para.runs:
        set_run_font(run, size=16, bold=True)
    return para


def h2(doc, text):
    para = doc.add_heading(text, level=2)
    for run in para.runs:
        set_run_font(run, size=14, bold=True)
    return para


def h3(doc, text):
    para = doc.add_heading(text, level=3)
    for run in para.runs:
        set_run_font(run, size=12, bold=True)
    return para


def bullets(doc, items):
    for it in items:
        para = doc.add_paragraph(it, style="List Bullet")
        for run in para.runs:
            set_run_font(run, size=12)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def numbered(doc, items):
    for it in items:
        para = doc.add_paragraph(it, style="List Number")
        for run in para.runs:
            set_run_font(run, size=12)


def add_table(doc, headers, rows, col_w=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                set_run_font(run, size=11, bold=True)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for para in cells[ci].paragraphs:
                for run in para.runs:
                    set_run_font(run, size=10)
    doc.add_paragraph()


def try_image(doc, path: Path, width_in=5.8, caption=None):
    if not path or not path.exists():
        p(doc, f"[Figure missing: {path}]", italic=True, size=10)
        return False
    try:
        doc.add_picture(str(path), width=Inches(width_in))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            p(doc, caption, size=10, italic=True, align="center", space_after=12)
        return True
    except Exception as e:
        p(doc, f"[Could not insert image {path.name}: {e}]", italic=True, size=10)
        return False


def find_shot(*names):
    for n in names:
        for base in (SHOT, LEGACY_SHOT, ROOT / "reports" / "screenshots"):
            pth = base / n
            if pth.exists():
                return pth
    # any jpeg in project_docs
    for pth in SHOT.glob("*.jpeg"):
        return pth
    for pth in SHOT.glob("*.jpg"):
        return pth
    for pth in SHOT.glob("*.png"):
        return pth
    return None


def build():
    doc = Document()
    # A4-ish margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run("CRIMECAST — Full Project Report  |  Page ")
        set_run_font(run, size=9)
        add_page_number(fp)
        run2 = fp.add_run("")
        set_run_font(run2, size=9)

    # ========== COVER ==========
    p(doc, "", space_after=40)
    p(doc, "A PROJECT REPORT", size=14, bold=True, align="center")
    p(doc, "ON", size=12, align="center")
    p(doc, "", space_after=6)
    p(
        doc,
        "CRIMECAST: Crime Analysis, Prediction and Sentiment Analysis\nUsing Machine Learning for Tamil Nadu Districts",
        size=16,
        bold=True,
        align="center",
    )
    p(doc, "", space_after=20)
    p(doc, "Submitted in partial fulfilment of the requirements for the award of the degree of", size=11, italic=True, align="center")
    p(doc, "[DEGREE NAME — e.g. B.E. / B.Tech / M.Sc.]", size=12, bold=True, align="center")
    p(doc, "in", size=11, align="center")
    p(doc, "[DEPARTMENT — e.g. Computer Science and Engineering]", size=12, bold=True, align="center")
    p(doc, "", space_after=24)
    p(doc, "Submitted by", size=11, align="center")
    p(doc, "[STUDENT NAME]", size=13, bold=True, align="center")
    p(doc, "Register No.: [REGISTER NUMBER]", size=12, align="center")
    p(doc, "", space_after=18)
    p(doc, "Under the guidance of", size=11, align="center")
    p(doc, "[GUIDE NAME, Designation]", size=12, bold=True, align="center")
    p(doc, "", space_after=28)
    p(doc, "[COLLEGE / UNIVERSITY NAME]", size=13, bold=True, align="center")
    p(doc, "[City, State]", size=11, align="center")
    p(doc, "Academic Year: 2025–2026", size=12, align="center")
    doc.add_page_break()

    # ========== CERTIFICATE / DECLARATION placeholders ==========
    h1(doc, "CERTIFICATE")
    p(
        doc,
        "This is to certify that the project report entitled “CRIMECAST: Crime Analysis, "
        "Prediction and Sentiment Analysis Using Machine Learning for Tamil Nadu Districts” "
        "submitted by [STUDENT NAME] (Register No. [NUMBER]) is a bonafide record of work "
        "carried out under my supervision. The contents of this report, in full or in parts, "
        "have not been submitted to any other Institute or University for the award of any "
        "degree or diploma.",
        align="justify",
    )
    p(doc, "", space_after=30)
    p(doc, "Signature of the Guide                          Signature of the HOD", size=11)
    p(doc, "", space_after=20)
    p(doc, "Place: _______________     Date: _______________", size=11)
    doc.add_page_break()

    h1(doc, "DECLARATION")
    p(
        doc,
        "I hereby declare that the project work entitled “CRIMECAST: Crime Analysis, Prediction "
        "and Sentiment Analysis Using Machine Learning for Tamil Nadu Districts” submitted to "
        "[College Name] is a record of original work done by me under the guidance of "
        "[Guide Name]. I further declare that this work has not been submitted elsewhere for "
        "any other degree or diploma.",
        align="justify",
    )
    p(doc, "", space_after=30)
    p(doc, "Signature of the Candidate", size=11)
    p(doc, "Name: [STUDENT NAME]", size=11)
    p(doc, "Register No.: [NUMBER]", size=11)
    doc.add_page_break()

    h1(doc, "ACKNOWLEDGEMENT")
    p(
        doc,
        "I express my sincere gratitude to my project guide [Guide Name] for continuous "
        "guidance and encouragement. I thank the Head of the Department and faculty members "
        "of [Department] for their support. I also acknowledge open data initiatives and "
        "public reporting that make academic study of crime statistics possible. Finally, "
        "I thank my family and friends for their support during the project work.",
        align="justify",
    )
    doc.add_page_break()

    # ========== INDEX (match screenshot) ==========
    h1(doc, "INDEX")
    add_table(
        doc,
        ["S.NO", "CONTENTS", "PAGE NO."],
        [
            ["01.", "INTRODUCTION", ""],
            ["", "1.1 COMPANY / ORGANIZATION PROFILE", ""],
            ["", "1.2 PROJECT OVERVIEW", ""],
            ["02.", "SYSTEM ANALYSIS", ""],
            ["", "2.1 FEASIBILITY STUDY", ""],
            ["", "2.2 EXISTING SYSTEM", ""],
            ["", "2.3 PROPOSED SYSTEM", ""],
            ["03.", "SYSTEM CONFIGURATION", ""],
            ["", "3.1 HARDWARE SPECIFICATION", ""],
            ["", "3.2 SOFTWARE SPECIFICATION", ""],
            ["", "3.3 ABOUT THE SOFTWARE", ""],
            ["04.", "SYSTEM DESIGN", ""],
            ["", "4.1 NORMALIZATION", ""],
            ["", "4.2 TABLE DESIGN", ""],
            ["", "4.3 INPUT DESIGN", ""],
            ["", "4.4 SFD / DFD", ""],
            ["05.", "SYSTEM DESCRIPTION", ""],
            ["06.", "TESTING AND IMPLEMENTATION", ""],
            ["07.", "CONCLUSION AND FUTURE SCOPE", ""],
            ["08.", "FORMS AND REPORT", ""],
            ["09.", "BIBLIOGRAPHY", ""],
        ],
    )
    p(
        doc,
        "Note: Page numbers can be filled after final print layout (or use Word Insert → Page Numbers / TOC).",
        size=10,
        italic=True,
    )
    doc.add_page_break()

    # ========== 01 INTRODUCTION ==========
    h1(doc, "01. INTRODUCTION")

    h2(doc, "1.1 COMPANY / ORGANIZATION PROFILE")
    p(
        doc,
        "This project is developed in an academic setting as a decision-support prototype for "
        "district-level crime analysis in Tamil Nadu. The “client domain” is public-safety "
        "analytics: analysts, students, and researchers who work with SCRB/NCRB-style tables "
        "and contemporaneous news coverage.",
        align="justify",
    )
    p(
        doc,
        "Official crime statistics are typically published with lag, split across multiple "
        "tables (complaints, crimes against women, murder/homicide), and use inconsistent "
        "headers and district labels. Media coverage provides timely signals but is not a "
        "substitute for FIRs. CRIMECAST integrates both layers carefully: official-era "
        "numeric labels train models; news and sentiment support live monitoring and explanation.",
        align="justify",
    )
    p(
        doc,
        "The system is a college prototype. It is not a live police system, not an official "
        "SCRB forecast product, and must not be used for automated enforcement decisions.",
        align="justify",
    )

    h2(doc, "1.2 PROJECT OVERVIEW")
    p(
        doc,
        "CRIMECAST (Crime Analysis, Prediction and Sentiment Analysis Using Machine Learning) "
        "is an end-to-end Python system with a Streamlit dashboard. It cleans multi-year Tamil "
        "Nadu district crime CSVs, trains regression models for counts and rates, scores news "
        "sentiment, produces 2026 scenario forecasts with uncertainty bands, and explains "
        "district risk using multi-source factors and model-based (SHAP-proxy / LIME-style) tools.",
        align="justify",
    )
    p(doc, "Major modules:", bold=True)
    bullets(
        doc,
        [
            "Data acquisition & cleaning (SCRB/NCRB-style CSVs → ML-ready table)",
            "Model training (Ridge / Random Forest / Gradient Boosting with log targets)",
            "Prediction with official-history blend for sticky rates",
            "News harvest (Tamil + English) and sentiment scoring (DistilBERT / lexicon)",
            "2026 multi-target forecasts (linear / last-year / blend) for TN38 districts",
            "Interactive dashboard: Live Feed, Map, Accuracy, Predict, Sentiment, Forecasts, Compare, Risk Explain, Health",
            "SQLite storage with optional CSV→DB migration (data/crimecast.db)",
        ],
    )
    p(doc, "Primary prediction targets:", bold=True)
    bullets(
        doc,
        [
            "Total complaints",
            "Murder incidence and murder rate",
            "Rape incidents (Sec. 376) and rape rate",
            "Cognizable crime rate (IPC+SLL)",
        ],
    )
    doc.add_page_break()

    # ========== 02 SYSTEM ANALYSIS ==========
    h1(doc, "02. SYSTEM ANALYSIS")

    h2(doc, "2.1 FEASIBILITY STUDY")
    h3(doc, "2.1.1 Technical Feasibility")
    p(
        doc,
        "The stack (Python, pandas, scikit-learn, Streamlit, Plotly, SQLite, optional "
        "transformers/DistilBERT) runs on a standard student laptop. Models are stored as "
        "joblib files; the dashboard needs no separate application server for demo. GeoJSON "
        "for Tamil Nadu districts is cached under assets/. Hence the project is technically feasible.",
        align="justify",
    )
    h3(doc, "2.1.2 Operational Feasibility")
    p(
        doc,
        "Users operate via a sidebar navigation dashboard. Batch scripts (START_DASHBOARD.bat, "
        "news refresh, health check, CSV migrate) support non-expert operation. Outputs (CSV, "
        "HTML district briefs) support report writing. Operationally feasible for academic demos.",
        align="justify",
    )
    h3(doc, "2.1.3 Economic Feasibility")
    p(
        doc,
        "Software used is free/open-source. No paid cloud API is required for core demos "
        "(news via public RSS-style harvest; sentiment has lexicon fallback). Economic cost is "
        "limited to student hardware time.",
        align="justify",
    )
    h3(doc, "2.1.4 Schedule Feasibility")
    p(
        doc,
        "The pipeline is modular: clean → train → visualize → dashboard. Incremental delivery "
        "allowed partial reports earlier and full integration for final submission.",
        align="justify",
    )

    h2(doc, "2.2 EXISTING SYSTEM")
    p(
        doc,
        "In conventional practice, analysts open separate spreadsheets for each crime category "
        "and year, clean them manually, and prepare charts in office tools. Problems include:",
        align="justify",
    )
    bullets(
        doc,
        [
            "Inconsistent column names and district spellings across years",
            "Risk of including aggregate “Total” rows (double counting)",
            "No systematic multi-algorithm model comparison",
            "Difficulty reproducing results after months",
            "News/sentiment disconnected from numeric crime tables",
            "No single interactive map for all 38 TN districts",
        ],
    )

    h2(doc, "2.3 PROPOSED SYSTEM")
    p(
        doc,
        "CRIMECAST automates discovery, cleaning, modelling, sentiment, forecasting, and "
        "visualization in one repository with a Streamlit UI.",
        align="justify",
    )
    p(doc, "Advantages of the proposed system:", bold=True)
    bullets(
        doc,
        [
            "Reproducible ML-ready table (crimecast_ml_ready.csv / SQLite ds_ml_ready)",
            "Official-year training labels (≤2023 era) vs media-proxy years for maps only",
            "Temporal validation for honest accuracy claims",
            "History-blended predictions to preserve rate rankings (e.g. sticky murder rates)",
            "Live news heat with explicit “not FIR” labelling",
            "TN38 rollup (city units → parent districts; junk units dropped)",
            "Explainability (composite risk factors + LIME-style / importance×z proxy)",
            "District Compare with carve maps, per-lakh metrics, history, HTML briefs",
        ],
    )
    doc.add_page_break()

    # ========== 03 SYSTEM CONFIGURATION ==========
    h1(doc, "03. SYSTEM CONFIGURATION")

    h2(doc, "3.1 HARDWARE SPECIFICATION")
    add_table(
        doc,
        ["Component", "Minimum / Recommended"],
        [
            ["Processor", "Intel i5 / AMD Ryzen 5 or better"],
            ["RAM", "8 GB minimum; 16 GB recommended (DistilBERT)"],
            ["Storage", "10 GB free (models, caches, news harvest)"],
            ["Display", "1366×768 minimum; 1920×1080 recommended for dashboard"],
            ["Network", "Optional (first GeoJSON download; news refresh)"],
        ],
    )

    h2(doc, "3.2 SOFTWARE SPECIFICATION")
    add_table(
        doc,
        ["Software", "Purpose"],
        [
            ["Windows 10/11 (or Linux/macOS)", "Host OS"],
            ["Python 3.10+", "Runtime"],
            ["pandas, numpy, scikit-learn", "Data & ML"],
            ["Streamlit, Plotly", "Interactive dashboard & charts"],
            ["joblib", "Model persistence"],
            ["SQLite (built-in)", "Local database data/crimecast.db"],
            ["transformers + torch (optional)", "DistilBERT sentiment"],
            ["wordcloud (optional)", "Sentiment word clouds"],
            ["Git (optional)", "Version control"],
        ],
    )

    h2(doc, "3.3 ABOUT THE SOFTWARE")
    p(
        doc,
        "Python is chosen for scientific computing ecosystem maturity. Streamlit allows rapid "
        "UI without separate frontend frameworks. scikit-learn provides transparent classical "
        "regressors suitable for tabular district data. Plotly produces interactive choropleths. "
        "SQLite stores headlines, sentiment aggregates, forecasts, alert logs, and migrated CSV "
        "datasets without a separate DB server.",
        align="justify",
    )
    p(doc, "How to run the dashboard:", bold=True)
    p(doc, "  cd CRIMECAST", size=11)
    p(doc, "  START_DASHBOARD.bat", size=11)
    p(doc, "  or:  py -3 -m streamlit run dashboard.py --server.port 8501", size=11)
    p(doc, "Browser: http://localhost:8501", size=11)
    doc.add_page_break()

    # ========== 04 SYSTEM DESIGN ==========
    h1(doc, "04. SYSTEM DESIGN")

    h2(doc, "4.1 NORMALIZATION")
    p(
        doc,
        "Although primary sources are denormalized yearly CSVs, the cleaned design moves toward "
        "analytical normal form suitable for ML:",
        align="justify",
    )
    bullets(
        doc,
        [
            "One observation grain: district_city × year (plus area_type flags)",
            "Separate fact families merged carefully (complaints, women crimes, murder)",
            "Removal of aggregate “Total / State” rows to avoid double counting",
            "Canonical district names via aliases and TN38 parent mapping (cities → districts)",
            "Rates stored separately from counts; population estimates fill missing lakhs",
            "Database tables split by concern: news_headlines, district_sentiment, rape_2026, alert_log, ds_* datasets",
        ],
    )
    p(
        doc,
        "SQLite dataset tables (ds_ml_ready, ds_media_harvest, …) hold full frames for dashboard "
        "loads; structured tables hold key operational entities. This is a hybrid warehouse style "
        "appropriate for a prototype (not a full 3NF enterprise ERP schema).",
        align="justify",
    )

    h2(doc, "4.2 TABLE DESIGN")
    p(doc, "A. File-based analytical tables", bold=True)
    add_table(
        doc,
        ["Table / File", "Key fields", "Role"],
        [
            ["crimecast_ml_ready.csv", "district_city, year, rates, counts, features", "ML training & maps"],
            ["fitted_predictions.csv", "district, year, target, actual, predicted", "Trend history"],
            ["training_metrics.csv", "target, model, MAE, R², temporal_*", "Model selection"],
            ["rape_predictions_2026_*.csv", "district, predicted, pred_low/high, risk", "Scenario forecast"],
            ["media_harvest_*.csv", "date, district, headline, source, lang", "News support"],
            ["sentiment_scores.csv", "text, polarity, label, crime_intensity", "Sentiment corpus"],
        ],
    )
    p(doc, "B. SQLite (data/crimecast.db) — structured", bold=True)
    add_table(
        doc,
        ["Table", "Purpose"],
        [
            ["meta", "Key-value sync timestamps"],
            ["news_headlines", "Harvested headlines + optional scores"],
            ["district_sentiment", "Aggregated polarity / concern by district"],
            ["rape_2026", "Forecast snapshot per district"],
            ["alert_log", "Persisted HIGH/MED operational alerts"],
            ["dataset_registry", "Catalogue of migrated CSV datasets"],
            ["ds_ml_ready, ds_media_harvest, …", "Full CSV frames in DB"],
        ],
    )

    h2(doc, "4.3 INPUT DESIGN")
    p(doc, "Inputs to the system:", bold=True)
    bullets(
        doc,
        [
            "Raw yearly CSVs under dataset/ (complaints, women crimes, murder/homicide)",
            "Optional SCRB/NCRB staged/drop-in files under dataset/scrb_ncrb/",
            "Dashboard controls: district, target, year, news time window, forecast method",
            "News refresh action (incremental harvest)",
            "Optional labeled sentiment text for TF-IDF training",
        ],
    )
    p(
        doc,
        "Input validation includes year parsing, numeric coercion, junk district filtering, "
        "and TN38 resolution so maps and scoreboards remain consistent.",
        align="justify",
    )

    h2(doc, "4.4 SFD / DFD")
    p(
        doc,
        "System Flow Diagram (SFD) and Data Flow Diagrams (DFD Level 0 / Level 1) document "
        "how data moves from sources through processes to users.",
        align="justify",
    )
    p(doc, "Figure 4.1 — Context diagram (DFD Level 0)", bold=True, align="center")
    try_image(doc, DIAG / "dfd_level_0.png", width_in=5.5, caption="Figure 4.1 DFD Level 0 — CRIMECAST context")
    p(doc, "Figure 4.2 — DFD Level 1", bold=True, align="center")
    try_image(doc, DIAG / "dfd_level_1.png", width_in=5.5, caption="Figure 4.2 DFD Level 1 — major processes")
    p(doc, "Figure 4.3 — System flow", bold=True, align="center")
    try_image(doc, DIAG / "system_flow_diagram.png", width_in=5.5, caption="Figure 4.3 System flow diagram")
    p(doc, "Level-0 summary:", bold=True)
    bullets(
        doc,
        [
            "External entities: Crime data sources; Text/news sources; User/Analyst",
            "Process 0: CRIMECAST System",
            "Outputs: Predictions, sentiment, 2026 scenarios, dashboard views, reports",
        ],
    )
    p(doc, "Level-1 processes (typical):", bold=True)
    bullets(
        doc,
        [
            "P1 Acquire / clean official tables",
            "P2 Feature engineering & ML-ready build",
            "P3 Train & evaluate models",
            "P4 Predict & blend with history",
            "P5 News harvest & sentiment",
            "P6 Forecast 2026 scenarios",
            "P7 Dashboard presentation & alerts",
        ],
    )
    doc.add_page_break()

    # ========== 05 SYSTEM DESCRIPTION ==========
    h1(doc, "05. SYSTEM DESCRIPTION")
    p(
        doc,
        "This chapter describes functional modules as implemented in the repository.",
        align="justify",
    )

    h3(doc, "5.1 Data Cleaning Module (clean_data.py)")
    p(
        doc,
        "Discovers supported year files, standardises headers, drops aggregate rows, coerces "
        "numerics, merges families into crimecast_ml_ready.csv. Flags distinguish official-era "
        "labels from thinner/proxy years used only as templates or map context.",
        align="justify",
    )

    h3(doc, "5.2 Training Module (train_model.py)")
    p(
        doc,
        "For each target, builds a preprocessing pipeline (numeric impute/scale + categorical "
        "encoding), compares Dummy, Ridge, Random Forest, and Gradient Boosting (often on log "
        "targets). Selection emphasises temporal holdout where available. Artifacts: models/*.joblib, "
        "best_models.json, training_metrics.csv, training_report.md.",
        align="justify",
    )
    p(doc, "Best models (from current training_report.md):", bold=True)
    add_table(
        doc,
        ["Target", "Best model", "Temporal MAE", "Temporal R²"],
        [
            ["Total complaints", "gradient_boosting_log", "44760.44", "-0.87"],
            ["Murder incidence", "gradient_boosting_log", "11.52", "0.44"],
            ["Rape incidents", "random_forest_log", "3.15", "0.36"],
            ["Murder rate", "gradient_boosting_log", "0.73", "0.50"],
            ["Rape rate", "ridge_log", "0.39", "0.56"],
            ["Cognizable rate", "gradient_boosting_log", "—", "—"],
        ],
    )
    p(
        doc,
        "Honest note: temporal R² for some count targets is weak with short history — the "
        "dashboard Accuracy tab exposes this rather than hiding it.",
        align="justify",
        italic=True,
    )

    h3(doc, "5.3 Prediction Module (predict.py)")
    p(
        doc,
        "Resolves districts (including TN38 parents), builds feature rows for a requested year, "
        "runs the stored pipeline, and blends model output with multi-year official history "
        "(stronger weight for rates). populate_all_district_predictions fills the TN map.",
        align="justify",
    )

    h3(doc, "5.4 Sentiment & News")
    p(
        doc,
        "News is harvested into media_harvest CSVs. Sentiment uses DistilBERT when available, "
        "else lexicon scoring. District concern scores drive the Sentiment map; word clouds "
        "visualise frequent terms per district.",
        align="justify",
    )

    h3(doc, "5.5 2026 Forecast Engine (forecast_engine.py / predict_2026_*)")
    p(
        doc,
        "District-level scenarios for rape incidents, murder incidence, and total complaints. "
        "Methods: linear trend, last-year carry, 50/50 blend. City units roll into TN38; "
        "uncertainty bands widen with horizon. Not official SCRB forecasts.",
        align="justify",
    )

    h3(doc, "5.6 Dashboard Modules (dashboard.py)")
    add_table(
        doc,
        ["Tab", "Function"],
        [
            ["Live Feed", "News heat, HIGH alerts, alert log, how-it-works, health strip"],
            ["District Map", "Choropleth, heat matrix, per-lakh scoreboard"],
            ["Accuracy Check", "Training metrics, claims, blend accuracy, holdout backtest"],
            ["Predict", "Single-area predict + all-district map (no news fill)"],
            ["Sentiment", "Concern map + word clouds"],
            ["2026 Forecasts", "Multi-target, method toggle, map, uncertainty"],
            ["District Compare", "Carve maps + full metrics + history + briefs"],
            ["Risk Explain", "Composite factors + SHAP-proxy + LIME-style"],
            ["Health", "File/model/DB checks; CSV→DB migrate"],
        ],
    )

    h3(doc, "5.7 Database Layer (db.py)")
    p(
        doc,
        "SQLite at data/crimecast.db. migrate_csv_to_db.py (or Health tab) loads main CSVs into "
        "ds_* tables. Dashboard loaders prefer DB then fall back to CSV. PostgreSQL is reserved "
        "for future multi-user deploy (env URL stub only).",
        align="justify",
    )
    doc.add_page_break()

    # ========== 06 TESTING ==========
    h1(doc, "06. TESTING AND IMPLEMENTATION")
    p(
        doc,
        "This chapter describes how CRIMECAST was implemented (environment and build steps) "
        "and how it was tested (strategy, formal cases, automated runner, UI checklist, and result figures).",
        align="justify",
    )

    h2(doc, "6.1 Implementation Environment")
    add_table(
        doc,
        ["Item", "Detail"],
        [
            ["Language", "Python 3.10+"],
            ["Core libraries", "pandas, numpy, scikit-learn, joblib"],
            ["UI", "Streamlit, Plotly"],
            ["Storage", "CSV files + SQLite (data/crimecast.db)"],
            ["Optional", "transformers/torch (DistilBERT), wordcloud"],
            ["OS (demo)", "Windows 10/11"],
            ["Test runner", "py -3 run_tests.py"],
            ["Dashboard", "py -3 -m streamlit run dashboard.py --server.port 8501"],
        ],
    )

    h2(doc, "6.2 Implementation Steps")
    numbered(
        doc,
        [
            "Data preparation — place yearly CSVs under dataset/ (optional SCRB staging).",
            "Cleaning — clean_data / pipeline → dataset/cleaned/crimecast_ml_ready.csv.",
            "Training — train_model.py; review training_report.md and training_metrics.csv.",
            "Prediction — predict.py with official-history blend for rates.",
            "News & sentiment — harvest refresh; DistilBERT/lexicon scoring; word clouds.",
            "2026 scenarios — forecast_engine / predict_2026 (linear · last-year · blend; TN38).",
            "Optional DB — migrate_csv_to_db.py or Health tab Migrate CSVs.",
            "Dashboard — streamlit run dashboard.py; walk Live → Map → Accuracy → Predict → 2026 → Compare.",
            "Exports — district brief HTML, accuracy CSV, figures, test terminal capture.",
        ],
    )

    h2(doc, "6.3 Testing Strategy")
    p(
        doc,
        "Testing is organised in layers P0–P4. Automated tests run without .bat files:",
        align="justify",
    )
    p(doc, "    py -3 run_tests.py", size=11)
    p(doc, "    py -3 project_docs/capture_test_terminal.py", size=11)
    add_table(
        doc,
        ["Layer", "Name", "Purpose"],
        [
            ["P0", "Core unit", "Districts, forecast math, layout, health structure"],
            ["P1", "Unit + fixtures", "clean_data, blend weights, alert rules, map aliases"],
            ["P2", "Integration", "predict_for_area; forecast_districts column contracts"],
            ["P3", "Data quality", "ML-ready / metrics / 2026 schema and non-negativity"],
            ["P4", "UI / system", "Manual checklist U1–U24; optional Streamlit AppTest"],
        ],
    )
    add_table(
        doc,
        ["Type", "What was checked in CRIMECAST"],
        [
            ["Unit testing", "Column normalise, TOTAL drop, TN38, forecast bands, blend, alerts"],
            ["Integration testing", "Predict + populate; 2026 multi-target schema"],
            ["Data-quality testing", "Year range, no aggregates, non-negative counts, official flags"],
            ["Model evaluation", "CV / test / temporal MAE and R² (Accuracy tab + metrics CSV)"],
            ["System / UI testing", "Pages load; forecast/predict maps without news fill"],
            ["Regression / self-test", "2026 engine remains sklearn-free"],
        ],
    )

    h2(doc, "6.4 Formal Test Cases (ID · Steps · Expected · Actual)")
    p(
        doc,
        "Eight formal test cases. Full narrative: docs/FORMAL_TEST_CASES.md. "
        "UI walkthrough: docs/MANUAL_UI_CHECKLIST.md (U1–U24).",
        align="justify",
    )
    add_table(
        doc,
        ["ID", "Type", "Steps (summary)", "Expected", "Actual / Status"],
        [
            ["TC-01", "Unit", "District entities (to_tn38, junk)", "Cities→parents; TN38=38", "Pass"],
            ["TC-02", "Unit", "clean_dataset drops TOTAL row", "TOTAL removed; year set", "Pass"],
            ["TC-03", "Unit", "Blend history rate vs count weights", "62%/35% hist; ≥0", "Pass"],
            ["TC-04", "Unit", "Alerts Thoothukudi>Madurai + 2026 HIGH", "HIGH alerts raised", "Pass"],
            ["TC-05", "Integration", "predict_for_area murder_rate Chennai", "prediction≥0", "Pass/Skip*"],
            ["TC-06", "Integration", "forecast_districts rape linear", "Schema + bands OK", "Pass"],
            ["TC-07", "Data quality", "ML-ready quality rules", "No TOTAL; non-neg", "Pass"],
            ["TC-08", "System/UI", "Dashboard + P4 source checks", "No crash; no news fill", "Pass"],
        ],
    )
    p(
        doc,
        "*TC-05 may Skip if trained .joblib models are absent; unit and data-quality tests still run.",
        size=10,
        italic=True,
    )
    p(doc, "Test execution: command  py -3 run_tests.py  ·  Overall Pass if failures=0 and errors=0.", size=11)
    p(doc, "Figure 6.0 — Terminal output of py -3 run_tests.py", bold=True, align="center")
    try_image(
        doc,
        pick_shot("shot_08_run_tests_terminal.png"),
        width_in=5.8,
        caption="Figure 6.0 Automated test run (run_tests.py) — testing evidence for Forms/Report",
    )
    p(
        doc,
        "Generate/update Figure 6.0:  py -3 project_docs/capture_test_terminal.py",
        size=10,
        italic=True,
        align="center",
    )

    h2(doc, "6.5 Implementation Modules Delivered")
    add_table(
        doc,
        ["Module", "File(s)", "Outcome"],
        [
            ["Cleaning", "clean_data.py", "Multi-year ML-ready table"],
            ["Training", "train_model.py", "Best model per target + metrics"],
            ["Prediction", "predict.py", "District predict + history blend"],
            ["Forecast 2026", "forecast_engine.py, predict_2026_*", "Scenarios + uncertainty"],
            ["Sentiment / news", "sentiment_*, harvest CSVs", "Concern map + word clouds"],
            ["Dashboard", "dashboard.py", "Live through Health tabs"],
            ["Database", "db.py, migrate_csv_to_db.py", "SQLite + CSV migrate"],
            ["Testing", "run_tests.py, tests/test_*.py", "P0–P4 automated suite"],
        ],
    )

    h2(doc, "6.6 Results Snapshot (current data — regenerated)")
    p(
        doc,
        "Figures below are regenerated from live project data into project_docs/figures/results/ "
        "(preferred over legacy model_outputs/figures). 2026 charts are scenarios, not SCRB facts; "
        "news volume is a media support layer, not FIRs.",
        align="justify",
        italic=True,
        size=11,
    )
    p(doc, "Figure 6.1 — Actual vs predicted (fitted models)", bold=True, align="center")
    try_image(
        doc,
        pick_fig("04_actual_vs_predicted.png", "actual_vs_predicted.png"),
        width_in=5.4,
        caption="Figure 6.1 Actual vs predicted (current fitted_predictions)",
    )
    p(doc, "Figure 6.2 — Top murder incidence (TN38, latest year)", bold=True, align="center")
    try_image(
        doc,
        pick_fig("01_top_murder_incidence.png", "top_murder_incidence.png"),
        width_in=5.2,
        caption="Figure 6.2 Top murder incidence — current ML-ready",
    )
    p(doc, "Figure 6.3 — Top rape incidents Sec.376 (latest)", bold=True, align="center")
    try_image(
        doc,
        pick_fig("02_top_rape_incidents.png", "top_rape_incidents.png"),
        width_in=5.2,
        caption="Figure 6.3 Top rape incidents — current ML-ready",
    )
    p(doc, "Figure 6.4 — Best models Test R²", bold=True, align="center")
    try_image(
        doc,
        pick_fig("05_training_test_r2.png"),
        width_in=5.2,
        caption="Figure 6.4 Training metrics — best models test R²",
    )
    p(doc, "Figure 6.5 — 2026 rape scenario top 15 (TN38)", bold=True, align="center")
    try_image(
        doc,
        pick_fig("06_rape_2026_top15.png", "rape_2026_top_districts.png"),
        width_in=5.2,
        caption="Figure 6.5 2026 scenario forecast — top districts (not official SCRB)",
    )
    p(doc, "Figure 6.6 — 2026 risk category share", bold=True, align="center")
    try_image(
        doc,
        pick_fig("07_rape_2026_risk_pie.png", "rape_2026_risk_pie.png"),
        width_in=4.2,
        caption="Figure 6.6 2026 risk categories",
    )
    p(doc, "Figure 6.7 — News volume by district (media harvest)", bold=True, align="center")
    try_image(
        doc,
        pick_fig("09_news_volume_by_district.png"),
        width_in=5.2,
        caption="Figure 6.7 Current news harvest volume (support layer, not FIRs)",
    )

    h2(doc, "6.7 Limitations of Testing")
    bullets(
        doc,
        [
            "Integration predict tests skip when models are not trained.",
            "Streamlit AppTest is optional and may skip in heavy (torch) environments.",
            "Manual UI checklist remains the primary UI sign-off for viva.",
            "Short official time series limits temporal R² — reported honestly on Accuracy page.",
        ],
    )
    doc.add_page_break()

    # ========== 07 CONCLUSION ==========
    h1(doc, "07. CONCLUSION AND FUTURE SCOPE")

    h2(doc, "7.1 Conclusion")
    p(
        doc,
        "CRIMECAST demonstrates a complete academic pipeline from messy multi-table crime CSVs "
        "to trained models, news-aware monitoring, scenario forecasts, and explainable district "
        "comparison. The design deliberately separates official training labels from media "
        "proxies, uses temporal evaluation, and presents uncertainty and limitations openly.",
        align="justify",
    )
    p(
        doc,
        "With limited official years, long-horizon forecasts remain scenarios. Within that "
        "constraint, the system is usable for teaching, demos, and structured analysis of "
        "Tamil Nadu district patterns.",
        align="justify",
    )

    h2(doc, "7.2 Limitations")
    bullets(
        doc,
        [
            "Short official time series limits true forecasting power",
            "Media headlines ≠ registered crimes",
            "Some administrative units require aggressive alias/junk filtering",
            "DistilBERT optional dependency can fail in constrained environments (lexicon fallback)",
            "PostgreSQL multi-user backend not fully implemented (SQLite is production path for demo)",
        ],
    )

    h2(doc, "7.3 Future Scope")
    bullets(
        doc,
        [
            "Add more official years (pre-2022 SCRB + consistent 2024+ when published)",
            "Monthly series for seasonality-aware models",
            "Full PostgreSQL adapter for multi-user lab deployment",
            "Stronger causal / policy evaluation designs (not only predictive)",
            "Multilingual Tamil NLP beyond lexicon for headlines",
            "User authentication and audit logs for institutional use",
        ],
    )
    doc.add_page_break()

    # ========== 08 FORMS AND REPORT ==========
    h1(doc, "08. FORMS AND REPORT")
    p(
        doc,
        "In a classical desktop system, “forms” are screens. In CRIMECAST, Streamlit pages "
        "serve as forms/reports. The following panels are regenerated from current project data "
        "(project_docs/figures/screenshots/) for the hard-copy annex — not legacy WhatsApp photos.",
        align="justify",
    )
    bullets(
        doc,
        [
            "Live Feed — news heat and alerts",
            "District Map — choropleth and scoreboard",
            "Accuracy Check — metrics and honesty statements",
            "2026 Forecasts — scenario map / ranking",
            "District Compare — multi-district analysis",
            "Sentiment — concern / volume panel",
            "Health — readiness checklist",
        ],
    )

    p(doc, "Figure 8.1 — Live Feed (form)", bold=True, align="center")
    try_image(doc, pick_shot("shot_01_live_feed.png"), width_in=5.6, caption="Figure 8.1 Form — Live Feed")
    p(doc, "Figure 8.2 — District Map & Scoreboard (form)", bold=True, align="center")
    try_image(doc, pick_shot("shot_02_district_map.png"), width_in=5.6, caption="Figure 8.2 Form — District Map")
    p(doc, "Figure 8.3 — Accuracy Check (form)", bold=True, align="center")
    try_image(doc, pick_shot("shot_03_accuracy.png"), width_in=5.6, caption="Figure 8.3 Form — Accuracy Check")
    p(doc, "Figure 8.4 — 2026 Forecasts (form)", bold=True, align="center")
    try_image(doc, pick_shot("shot_04_forecast_2026.png"), width_in=5.6, caption="Figure 8.4 Form — 2026 Forecasts")
    p(doc, "Figure 8.5 — District Compare (form)", bold=True, align="center")
    try_image(doc, pick_shot("shot_05_district_compare.png"), width_in=5.6, caption="Figure 8.5 Form — District Compare")
    p(doc, "Figure 8.6 — Sentiment (form)", bold=True, align="center")
    try_image(doc, pick_shot("shot_06_sentiment.png"), width_in=5.4, caption="Figure 8.6 Form — Sentiment")
    p(doc, "Figure 8.7 — Health (form)", bold=True, align="center")
    try_image(doc, pick_shot("shot_07_health.png"), width_in=5.4, caption="Figure 8.7 Form — Health check")

    p(doc, "Figure 8.8 — run_tests.py terminal (testing evidence)", bold=True, align="center")
    try_image(
        doc,
        pick_shot("shot_08_run_tests_terminal.png"),
        width_in=5.6,
        caption="Figure 8.8 Terminal screenshot — automated unit/integration tests",
    )
    p(
        doc,
        "Generate/update this figure:  py -3 project_docs/capture_test_terminal.py",
        size=10,
        italic=True,
        align="center",
    )

    p(doc, "Figure 8.9 — Report INDEX template (college structure)", bold=True, align="center")
    idx_img = find_shot(
        "WhatsApp Image 2026-07-19 at 2.58.51 PM.jpeg",
        "WhatsApp Image 2026-07-12 at 2.15.37 PM.jpeg",
    )
    try_image(doc, idx_img, width_in=4.2, caption="Figure 8.9 College INDEX structure (reference only)")

    p(
        doc,
        "Optional: after running the live Streamlit app, replace any panel above with a true "
        "browser screenshot (Win+Shift+S) for maximum realism in the bound book.",
        size=10,
        italic=True,
    )
    p(doc, "Printed annex suggestions:", bold=True)
    bullets(
        doc,
        [
            "training_report.md (print or paste metrics table)",
            "rape_predictions_2026_report.txt",
            "One district brief HTML printed via browser Ctrl+P",
            "health_check output",
            "project_docs/figures/MANIFEST.md (list of regenerated assets)",
        ],
    )
    doc.add_page_break()

    # ========== 09 BIBLIOGRAPHY ==========
    h1(doc, "09. BIBLIOGRAPHY")
    refs = [
        "National Crime Records Bureau (NCRB), Ministry of Home Affairs, Government of India — Crime in India (various years).",
        "State Crime Records Bureau (SCRB) / Tamil Nadu Police — district statistical tables (as available).",
        "OpenCity / public TN crime data portals — district CSV releases (e.g. 2023 sets used in academic workflows).",
        "Pedregosa et al., “Scikit-learn: Machine Learning in Python,” JMLR, 2011.",
        "McKinney, W., “Data Structures for Statistical Computing in Python” (pandas), Proc. of the 9th Python in Science Conference, 2010.",
        "Streamlit Inc., Streamlit Documentation — https://docs.streamlit.io/",
        "Plotly Technologies, Plotly.py Documentation — https://plotly.com/python/",
        "Sanh et al., “DistilBERT, a distilled version of BERT,” arXiv:1910.01108, 2019.",
        "Ribeiro, Singh, Guestrin, “Why Should I Trust You?” Explaining the Predictions of Any Classifier (LIME), KDD, 2016.",
        "Lundberg & Lee, “A Unified Approach to Interpreting Model Predictions” (SHAP), NeurIPS, 2017.",
        "Python Software Foundation, Python Language Reference — https://docs.python.org/3/",
        "SQLite Development Team, SQLite Documentation — https://www.sqlite.org/docs.html",
    ]
    for i, r in enumerate(refs, 1):
        p(doc, f"[{i}]  {r}", size=11, space_after=6)

    p(doc, "", space_after=20)
    p(doc, "— End of Report —", bold=True, align="center")
    p(
        doc,
        "Fill bracketed fields [STUDENT NAME], [REGISTER NUMBER], [GUIDE], [COLLEGE] before final print.",
        size=10,
        italic=True,
        align="center",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"[OK] Wrote {OUT}")
    return OUT


if __name__ == "__main__":
    build()
